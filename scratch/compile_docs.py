import os
import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# XML manipulation helpers for styling
def set_cell_background(cell, hex_color):
    """Sets background color (shading) for table cells."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_number(run):
    """Inserts a dynamic PAGE number field inside a header/footer run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_header_footer(doc):
    """Configures page headers and footers with project titles and dynamic page numbers."""
    section = doc.sections[0]
    
    # Configure Header
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "" # Clear default
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header_p.add_run("SOLIX Platform — Premium Technical Documentation")
    run_h.font.name = "Segoe UI"
    run_h.font.size = Pt(8.5)
    run_h.font.color.rgb = RGBColor(128, 128, 128)
    
    # Configure Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "" # Clear default
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run("Page ")
    run_f.font.name = "Segoe UI"
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(128, 128, 128)
    
    # Insert dynamic page field
    page_run = footer_p.add_run()
    page_run.font.name = "Segoe UI"
    add_page_number(page_run)
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = RGBColor(128, 128, 128)

def format_bold_text(paragraph, text, font_name="Segoe UI"):
    """
    Parses Markdown bold syntax (**text**) and appends formatted runs to a paragraph.
    Handles standard text and bold runs.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.font.name = font_name
        else:
            if part:
                run = paragraph.add_run(part)
                run.font.name = font_name

def build_cover_page(doc):
    """Creates a visually stunning, professional cover page for the document."""
    # Large spacing at the top
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Enterprise Technical Documentation\nSOLIX Virtual Data Engineer")
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(28)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(11, 25, 44) # Deep Navy
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("A Comprehensive Technical and Architectural Specification of the SOL Data Agent")
    run_sub.font.name = "Segoe UI"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(30, 62, 98) # Secondary Navy
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Institution / Team details
    p_team = doc.add_paragraph()
    p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_team_label = p_team.add_run("Prepared & Developed by:\n")
    run_team_label.font.name = "Segoe UI"
    run_team_label.font.size = Pt(11)
    run_team_label.font.color.rgb = RGBColor(128, 128, 128)
    
    run_team_name = p_team.add_run("The SOL Team (Space Of Learning)\n")
    run_team_name.font.name = "Segoe UI"
    run_team_name.font.size = Pt(15)
    run_team_name.bold = True
    run_team_name.font.color.rgb = RGBColor(11, 25, 44)
    
    run_hitu = p_team.add_run("Helwan International Technological University (HITU)\nCairo Technological College\n")
    run_hitu.font.name = "Segoe UI"
    run_hitu.font.size = Pt(11)
    run_hitu.font.color.rgb = RGBColor(44, 62, 80)
    
    for _ in range(3):
        doc.add_paragraph()
        
    # Supervisors
    p_sup = doc.add_paragraph()
    p_sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sup = p_sup.add_run("Under the Supervision of:\n")
    run_sup.font.name = "Segoe UI"
    run_sup.font.size = Pt(11)
    run_sup.font.color.rgb = RGBColor(128, 128, 128)
    
    run_names = p_sup.add_run("Academic Supervisor: Dr. Simon Ezzat\nCo-Supervisor: Eng. Naglaa Saeed")
    run_names.font.name = "Segoe UI"
    run_names.font.size = Pt(12)
    run_names.bold = True
    run_names.font.color.rgb = RGBColor(30, 62, 98)
    
    # Add page break
    doc.add_page_break()

def build_table_of_contents(doc):
    """Creates a beautifully styled, professional Table of Contents using dotted leaders."""
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(18)
    
    run_title = p_title.add_run("Table of Contents")
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(20)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(11, 25, 44)
    
    # Add divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    run_div = p_div.add_run("—" * 65)
    run_div.font.color.rgb = RGBColor(200, 200, 200)
    
    # Table of contents items (Name, Page number, Level)
    # Level 1 = Chapter, Level 2 = Section, Level 3 = Subsection
    toc_items = [
        ("Chapter 1: SOL Team and General Project Overview", "3", 1),
        ("Section 1: The SOL (Space Of Learning) Team", "3", 2),
        ("1.1 Who We Are: Vision and Mission of the SOL Team", "3", 3),
        ("1.2 Documented Awards and Hackathon Achievements", "3", 3),
        ("1.3 Academic and Institutional Supervision", "4", 3),
        ("1.4 Leadership and Co-Founders of SOL", "4", 3),
        ("1.5 Roles and Responsibilities of the 16 Team Members", "4", 3),
        ("Section 2: General Introduction to SOL Data Agent (SOLIX)", "6", 2),
        ("2.1 The General Project Concept", "6", 3),
        ("2.2 Problem Statement & Proposed Solution", "6", 3),
        ("2.3 Technical and Business Goals", "7", 3),
        ("2.4 High-Level Feature Catalog", "7", 3),
        
        ("Chapter 2: System Architecture and Directory Mapping", "8", 1),
        ("Section 1: Software Architecture and System Design", "8", 2),
        ("1.1 Architectural Overview (Enterprise Data Factory Model)", "8", 2),
        ("1.2 Data Flow and Processing Lifecycle", "9", 2),
        ("Section 2: Directory Structure and Codebase Mapping", "10", 2),
        ("2.1 Project Directory Tree and File Roles", "10", 2),
        ("2.2 Deep Dive: Core Code Module Analysis", "11", 2),
        ("1. The Sandbox Engine (core/copilot/sandbox.py)", "11", 3),
        ("2. The Universal Data Loader (data_layer/loaders/universal_loader.py)", "11", 3),
        ("3. Database ORM Schemas (backend/models.py)", "12", 3),
        
        ("Chapter 3: Technical and Functional Feature Catalog", "13", 1),
        ("Section 1: Universal Data Loader", "13", 2),
        ("1.1 Supported Data Formats", "13", 3),
        ("1.2 Ingestion Pipeline & Factory Registry", "13", 3),
        ("1.3 Memory Optimization & Preview Cache", "14", 3),
        ("1.4 Active Caching and Registry Map", "14", 3),
        ("Section 2: AI Dataset Advisor", "14", 2),
        ("2.1 Technical Architecture", "14", 3),
        ("2.2 Semantic Intent Parser", "15", 3),
        ("2.3 Local & Remote Vector Indexing", "15", 3),
        ("2.4 Hybrid Scoring and Reranking Layer", "15", 3),
        ("Section 3: Synthetic Data Studio", "16", 2),
        ("3.1 Kaggle Cloud Orchestrator", "16", 3),
        ("3.2 Mock Generator and Fallback Synthesis", "16", 3),
        ("Section 4: Chaos Corruptor Engine", "17", 2),
        ("4.1 Injection Anomaly Types", "17", 3),
        ("Section 5: OCR Data Extraction Engine", "18", 2),
        ("5.1 Architecture & Native Dependencies", "18", 3),
        ("5.2 PDF Processing Pipeline", "18", 3),
        ("5.3 Tabular Structure Parsing (text_to_dataframe)", "19", 3),
        
        ("Chapter 4: SOL Voice Copilot Engine", "20", 1),
        ("Section 1: Voice Copilot Architecture", "20", 2),
        ("Section 2: Agent 1 - STT Post-Processor & Intent Router", "20", 2),
        ("2.1 Intent Classification", "20", 3),
        ("2.2 Arabized Data Slang Mapping", "21", 3),
        ("Section 3: ReAct Thinking & Tool Execution Loop", "21", 2),
        ("3.1 ReAct Loop Details", "21", 3),
        ("Section 4: Secure Code Execution Sandbox", "22", 2),
        ("4.1 AST Safety Verification (_SecurityVisitor)", "22", 3),
        ("4.2 Namespace Isolation", "23", 3),
        ("4.3 Output and Exception Capturing", "23", 3),
        ("Section 5: Response Formatter & Egyptian Arabic Voice Synthesis", "24", 2),
        ("5.1 Persona and Voice Laws", "24", 3),
        ("5.2 ElevenLabs & Fallback Audio Pipeline", "25", 3),
        
        ("Chapter 5: Backend Development & API Core Specification", "26", 1),
        ("Section 1: FastAPI Core Architecture & Configurations", "26", 2),
        ("1.1 CORSMiddleware & Session Configurations", "26", 3),
        ("Section 2: SimpleRateLimiter Middleware", "27", 2),
        ("2.1 Mechanics of the Rate Limiter", "27", 3),
        ("Section 3: Authentication, Authorization, and Security Flows", "28", 2),
        ("3.1 Secure User Registration and Password Strength Validation", "28", 3),
        ("3.2 OTP Generation, Hashing, and Lockout Checks", "29", 3),
        ("3.3 JWT Session Tokens & Cookie Policies", "29", 3),
        ("3.4 Google and GitHub OAuth Integrations", "30", 3),
        ("Section 4: REST API Endpoints Specification", "31", 2),
        ("4.1 Undo and Redo Mechanics", "31", 3),
        ("Section 5: CleaningStudioPDFReportGenerator", "32", 2),
        ("5.1 Multilingual Support (Arabic & English RTL Rendering)", "32", 3),
        ("5.2 Matplotlib Visual Embedding", "33", 3),
        ("5.3 Compiled Sections", "33", 3),
        ("Section 6: Expected Defense Committee Questions and Answers", "34", 2),
        
        ("Chapter 6: Database Design & Schema Specification", "35", 1),
        ("Section 1: Relational Database Architecture Overview", "35", 2),
        ("Section 2: Platform Core Database Schema (sol.db)", "36", 2),
        ("2.1 Table: users", "36", 3),
        ("2.2 Table: otp_sessions", "37", 3),
        ("2.3 Table: auth_logs", "37", 3),
        ("2.4 Table: job_records", "38", 3),
        ("2.5 Table: projects", "38", 3),
        ("2.6 Table: tasks", "39", 3),
        ("2.7 Table: task_runs", "39", 3),
        ("2.8 Tables: forms & responses", "40", 3),
        ("2.9 Tables: notifications, token_usage_records & feedbacks", "40", 3),
        ("Section 3: Dataset Advisor Database Schema (advisor.db)", "41", 2),
        ("3.1 Table: datasets", "41", 3),
        ("3.2 Table: search_logs", "42", 3),
        ("3.3 Table: recommendations", "42", 3),
        ("Section 4: Schema Migrations and Maintenance", "43", 2),
        ("Section 5: Vector Search Database (Qdrant DB Setup)", "44", 2),
        ("Section 6: Expected Defense Committee Questions and Answers", "45", 2),
        
        ("Chapter 7: DevOps, Local Deployment & Testing Suite", "46", 1),
        ("Section 1: Local Deployment & System Requirements", "46", 2),
        ("1.1 System Requirements", "46", 3),
        ("1.2 Automated Startup Script (run_project.bat)", "47", 3),
        ("Section 2: Environment Variable Configurations (.env)", "48", 2),
        ("Section 3: Testing Framework (Pytest Suite)", "49", 2),
        ("3.1 Pytest Architecture & Dependency Overrides", "49", 3),
        ("3.2 Mocking Strategies & Sandbox Verification", "50", 3),
        ("Section 4: Expected Defense Committee Questions and Answers", "51", 2),
        
        ("Chapter 8: Appendix & Expected Graduation Committee Questions", "52", 1),
        ("Section 1: Technical Glossary", "52", 2),
        ("Section 2: Expected Oral Defense Questions and Model Answers", "53", 2),
    ]
    
    for title, page, level in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        
        # Configure tab stops for dotted leaders
        tab_stops = p.paragraph_format.tab_stops
        # Add tab stop at 6.0 inches, aligned right with dots leader
        tab_stops.add_tab_stop(Inches(6.0), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        
        # Adjust indents based on level
        if level == 1:
            p.paragraph_format.left_indent = Inches(0.0)
            run = p.add_run(title)
            run.font.name = "Segoe UI"
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(11, 25, 44)
            
            run_page = p.add_run(f"\t{page}")
            run_page.font.name = "Segoe UI"
            run_page.bold = True
            run_page.font.size = Pt(11)
            run_page.font.color.rgb = RGBColor(11, 25, 44)
        elif level == 2:
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(title)
            run.font.name = "Segoe UI"
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(30, 62, 98)
            
            run_page = p.add_run(f"\t{page}")
            run_page.font.name = "Segoe UI"
            run_page.bold = True
            run_page.font.size = Pt(10)
            run_page.font.color.rgb = RGBColor(30, 62, 98)
        else:
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(title)
            run.font.name = "Segoe UI"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(44, 62, 80)
            
            run_page = p.add_run(f"\t{page}")
            run_page.font.name = "Segoe UI"
            run_page.font.size = Pt(9.5)
            run_page.font.color.rgb = RGBColor(44, 62, 80)
            
    # Add page break
    doc.add_page_break()

def parse_markdown_to_docx(markdown_path, doc):
    """
    Parses a Markdown file and appends the content into the Word document
    using styled paragraphs, tables, lists, headings, and code blocks.
    """
    if not os.path.exists(markdown_path):
        print(f"File {markdown_path} not found.")
        return
        
    with open(markdown_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 1. Code Block Handling
        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                # Write code block to docx
                p_code = doc.add_paragraph()
                p_code.paragraph_format.left_indent = Inches(0.4)
                p_code.paragraph_format.right_indent = Inches(0.4)
                p_code.paragraph_format.space_before = Pt(6)
                p_code.paragraph_format.space_after = Pt(6)
                p_code.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                run_code = p_code.add_run("\n".join(code_content))
                run_code.font.name = "Consolas"
                run_code.font.size = Pt(9.5)
                run_code.font.color.rgb = RGBColor(44, 62, 80)
                
                # Set a light gray background for code block paragraphs
                pPr = p_code._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F4F7')
                pPr.append(shd)
                
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line.rstrip('\n'))
            i += 1
            continue
            
        # 2. Table Handling
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            in_table = False
            render_docx_table(table_rows, doc)
            table_rows = []
            if not stripped:
                i += 1
                continue
                
        # 3. Heading Handling
        if stripped.startswith("#"):
            match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if match:
                level = len(match.group(1))
                title = match.group(2)
                
                if level == 1:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(24)
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.keep_with_next = True
                    
                    run = p.add_run(title)
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(18)
                    run.bold = True
                    run.font.color.rgb = RGBColor(11, 25, 44) # Deep Navy
                    
                elif level == 2:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(18)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.keep_with_next = True
                    
                    run = p.add_run(title)
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(14)
                    run.bold = True
                    run.font.color.rgb = RGBColor(30, 62, 98) # Secondary Navy
                    
                elif level == 3:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.keep_with_next = True
                    
                    run = p.add_run(title)
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(11.5)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 141, 218) # Sky Blue
                    
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(title)
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(11)
                    run.bold = True
                    run.font.color.rgb = RGBColor(44, 62, 80)
            i += 1
            continue
            
        # 4. List Items Handling
        if stripped.startswith("* ") or stripped.startswith("- "):
            list_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.25)
            format_bold_text(p, list_text)
            i += 1
            continue
            
        match_numbered = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if match_numbered:
            num = match_numbered.group(1)
            list_text = match_numbered.group(2)
            p = doc.add_paragraph(style='List Number')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.25)
            format_bold_text(p, list_text)
            i += 1
            continue
            
        # 5. Horizontal rule divider
        if stripped == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("❖ ❖ ❖")
            run.font.name = "Segoe UI"
            run.font.color.rgb = RGBColor(128, 128, 128)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
            
        # 6. Standard Paragraph Handling
        if stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            format_bold_text(p, stripped)
                
        i += 1

def render_docx_table(table_rows, doc):
    """
    Renders a Markdown table into a Word Table with custom formatting:
    Deep navy header, light-gray alternative row borders, justified cell text, and LTR alignment.
    """
    if not table_rows or len(table_rows) < 2:
        return
        
    # Split cells and filter empty strings from splitting edges
    headers = [c.strip() for c in table_rows[0].split("|")[1:-1]]
    num_cols = len(headers)
    
    # Verify separator row
    separator_idx = 1
    if separator_idx < len(table_rows) and "---" in table_rows[separator_idx]:
        data_rows_raw = table_rows[separator_idx+1:]
    else:
        data_rows_raw = table_rows[separator_idx:]
        
    data_rows = []
    for r in data_rows_raw:
        cells = [c.strip() for c in r.split("|")[1:-1]]
        # Match column count
        if len(cells) < num_cols:
            cells.extend([""] * (num_cols - len(cells)))
        elif len(cells) > num_cols:
            cells = cells[:num_cols]
        data_rows.append(cells)
        
    # Create Word Table
    table = doc.add_table(rows=len(data_rows) + 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 1. Format Header Row
    hdr_cells = table.rows[0].cells
    for col_idx in range(num_cols):
        hdr_cells[col_idx].text = "" # Clear default
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        run = p.add_run(headers[col_idx])
        run.font.name = "Segoe UI"
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # White text
        run.font.size = Pt(10)
        
        set_cell_background(hdr_cells[col_idx], "0B192C") # Deep Navy
        
    # 2. Format Data Rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F8F9FA" if row_idx % 2 == 0 else "FFFFFF" # Shading
        
        for col_idx in range(num_cols):
            row_cells[col_idx].text = "" # Clear default
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            
            # Replace HTML line breaks with newlines before formatting
            cell_text = row_data[col_idx].replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
            format_bold_text(p, cell_text)
            
            # Apply consistent font size and name to all runs in the cell
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = "Segoe UI"
                
            set_cell_background(row_cells[col_idx], bg_color)
            
    # Add a spacing paragraph after the table
    doc.add_paragraph()

if __name__ == "__main__":
    print("Initializing document...")
    doc = Document()
    
    # Page Margins Configuration (Standard 1 Inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Build Cover Page
    print("Building cover page...")
    build_cover_page(doc)
    
    # Build Table of Contents
    print("Building table of contents...")
    build_table_of_contents(doc)
    
    # Configure Headers & Footers
    print("Configuring headers and footers...")
    add_header_footer(doc)
    
    # Parse Chapter 1
    print("Parsing Chapter 1...")
    chapter1_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter1.md"
    parse_markdown_to_docx(chapter1_path, doc)
    
    # Page Break between chapters
    doc.add_page_break()
    
    # Parse Chapter 2
    print("Parsing Chapter 2...")
    chapter2_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter2.md"
    parse_markdown_to_docx(chapter2_path, doc)
    
    # Page Break between chapters
    doc.add_page_break()
    
    # Parse Chapter 3
    print("Parsing Chapter 3...")
    chapter3_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter3.md"
    parse_markdown_to_docx(chapter3_path, doc)
    
    # Page Break between chapters
    doc.add_page_break()
    
    # Parse Chapter 4
    print("Parsing Chapter 4...")
    chapter4_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter4.md"
    parse_markdown_to_docx(chapter4_path, doc)
    
    # Page Break between chapters
    doc.add_page_break()
    
    # Parse Chapter 5
    print("Parsing Chapter 5...")
    chapter5_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter5.md"
    parse_markdown_to_docx(chapter5_path, doc)
    
    # Page Break between chapters
    doc.add_page_break()
    
    # Parse Chapter 6
    print("Parsing Chapter 6...")
    chapter6_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter6.md"
    parse_markdown_to_docx(chapter6_path, doc)
    
    # Page Break between chapters
    doc.add_page_break()
    
    # Parse Chapter 7
    print("Parsing Chapter 7...")
    chapter7_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter7.md"
    parse_markdown_to_docx(chapter7_path, doc)
    
    # Page Break between chapters
    doc.add_page_break()
    
    # Parse Chapter 8
    print("Parsing Chapter 8...")
    chapter8_path = r"e:\run-20260221T125607Z-1-001\run\scratch\chapter8.md"
    parse_markdown_to_docx(chapter8_path, doc)
    
    # Save the output file in the root directory
    output_filename = r"e:\run-20260221T125607Z-1-001\run\SOLIX_Project_Premium_Documentation.docx"
    print(f"Saving file to {output_filename}...")
    try:
        doc.save(output_filename)
        print("Documentation compiled successfully!")
    except PermissionError:
        fallback_filename = r"e:\run-20260221T125607Z-1-001\run\SOLIX_Project_Premium_Documentation_Final.docx"
        print(f"[WARNING] Permission denied on {output_filename} (likely open in MS Word).")
        print(f"Attempting fallback to: {fallback_filename}")
        doc.save(fallback_filename)
        print("Documentation compiled successfully as Final fallback!")
